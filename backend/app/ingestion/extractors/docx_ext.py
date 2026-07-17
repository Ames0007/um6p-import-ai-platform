"""Extraction DOCX (python-docx). Une « page » logique = le document entier."""
from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

from app.ingestion.extractors.base import CorruptedDocumentError, Extractor
from app.ingestion.types import ExtractedPage


class DocxExtractor(Extractor):
    extensions = {".docx"}

    def _open(self, path: str | Path):
        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx n'est pas installé.") from exc
        try:
            return docx.Document(str(path))
        except Exception as exc:
            raise CorruptedDocumentError(f"DOCX illisible : {exc}") from exc

    def count_pages(self, path: str | Path) -> int:
        # Le nombre de pages « rendues » n'est pas connu sans mise en page ;
        # on considère une unité logique.
        return 1

    def iter_pages(self, path: str | Path) -> Iterator[ExtractedPage]:
        document = self._open(path)
        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        yield ExtractedPage(number=1, text=text)
